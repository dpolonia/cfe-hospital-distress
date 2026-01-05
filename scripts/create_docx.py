"""
Convert V6 paper to DOCX format using python-docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re

def create_v6_docx():
    """Create V6 DOCX from content."""
    
    doc = Document()
    
    # Set document styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    
    # Add some space at top
    for _ in range(3):
        doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Vertical Integration and Financial Distress in Portuguese Public Hospitals")
    run.bold = True
    run.font.size = Pt(16)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Methodological Framework and Preliminary Evidence from the 2024 ULS Reform")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Author
    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.add_run("Daniel Ferreira Polónia").bold = True
    
    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil.add_run("Department of Economics, Management, Industrial Engineering and Tourism (DEGEIT)")
    
    affil2 = doc.add_paragraph()
    affil2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil2.add_run("University of Aveiro, Portugal")
    
    doc.add_paragraph()
    
    email = doc.add_paragraph()
    email.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    email.add_run("Corresponding Author: dpolonia@ua.pt")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadata box
    meta = doc.add_paragraph()
    meta.add_run("Submitted to: ").bold = True
    meta.add_run("Health Economics (Wiley)")
    
    meta2 = doc.add_paragraph()
    meta2.add_run("Running Head: ").bold = True
    meta2.add_run("ULS Reform: Framework and Preliminary Evidence")
    
    meta3 = doc.add_paragraph()
    meta3.add_run("Word Count: ").bold = True
    meta3.add_run("9,847 (excluding tables, figures, and references)")
    
    meta4 = doc.add_paragraph()
    meta4.add_run("Keywords: ").bold = True
    meta4.add_run("Health care reform; Hospital finance; Vertical integration; Difference-in-differences; Synthetic control; Portugal; Financial distress")
    
    meta5 = doc.add_paragraph()
    meta5.add_run("JEL Codes: ").bold = True
    meta5.add_run("I18, H51, G33, L22, C23")
    
    doc.add_page_break()
    
    # =========================================================================
    # ABSTRACT
    # =========================================================================
    
    doc.add_heading("Abstract", level=1)
    
    obj = doc.add_paragraph()
    obj.add_run("Objectives: ").bold = True
    obj.add_run("To establish a methodological framework for evaluating Portugal's 2024 universal healthcare integration reform and provide preliminary evidence on short-term financial trajectories.")
    
    meth = doc.add_paragraph()
    meth.add_run("Methods: ").bold = True
    meth.add_run("We compare 31 newly-integrated ULS entities against 8 pre-existing ULS (integrated 1999–2012) using difference-in-differences, synthetic control, and triple-difference designs with quarterly panel data from 2014–2025 (n=1,092 entity-quarters). The 3 IPO oncology centers are excluded from analysis.")
    
    res = doc.add_paragraph()
    res.add_run("Results: ").bold = True
    res.add_run("The primary DiD estimate is +6.2 percentage points (95% CI: −9.3 to +21.7; p=0.433). These estimates are severely underpowered (25% power; MDE=15pp). Pre-existing ULS exhibited 18.5pp higher baseline distress (p=0.006).")
    
    conc = doc.add_paragraph()
    conc.add_run("Conclusions: ").bold = True
    conc.add_run("This analysis cannot determine whether Portugal's 2024 reform improved, worsened, or had no effect on hospital financial distress. The paper's primary contribution is methodological.")
    
    doc.add_page_break()
    
    # =========================================================================
    # INTRODUCTION
    # =========================================================================
    
    doc.add_heading("1. Introduction", level=1)
    
    doc.add_paragraph(
        "Healthcare system integration—the consolidation of hospital and primary care services under unified management—has become a central policy lever for attempting to improve efficiency and coordination across healthcare systems internationally. In January 2024, Portugal implemented one of Europe's most comprehensive integration reforms, merging all public hospitals with primary care clusters (Agrupamentos de Centros de Saúde, ACES) into integrated Unidades Locais de Saúde (ULS; Local Health Units)."
    )
    
    doc.add_heading("1.1 The Identification Problem", level=2)
    
    doc.add_paragraph(
        "The core identification problem is this: Portugal's comparison group consists not of 'untreated' entities but of entities that experienced integration 12–25 years ago. When we compare newly-integrated entities against pre-existing ULS, we estimate the difference between short-run integration effects (among new adopters) and long-run integration effects (among early adopters)."
    )
    
    doc.add_heading("1.2 Contribution and Scope", level=2)
    
    doc.add_paragraph("This paper makes four contributions:")
    
    contrib = doc.add_paragraph(style='List Number')
    contrib.add_run("Methodological framework: ").bold = True
    contrib.add_run("Baseline methods for ongoing reform evaluation")
    
    contrib2 = doc.add_paragraph(style='List Number')
    contrib2.add_run("Preliminary descriptive evidence: ").bold = True
    contrib2.add_run("Initial point estimates establishing baseline magnitudes")
    
    contrib3 = doc.add_paragraph(style='List Number')
    contrib3.add_run("Selection pattern documentation: ").bold = True
    contrib3.add_run("Quantifying selection pattern in early ULS adoption")
    
    contrib4 = doc.add_paragraph(style='List Number')
    contrib4.add_run("Heterogeneity patterns: ").bold = True
    contrib4.add_run("Entity characteristics associated with larger effects")
    
    doc.add_page_break()
    
    # =========================================================================
    # INSTITUTIONAL CONTEXT
    # =========================================================================
    
    doc.add_heading("3. Institutional Context", level=1)
    
    doc.add_heading("3.1 Current SNS Structure (2025)", level=2)
    
    # Table 1: SNS Structure
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Entity Type"
    hdr[1].text = "Description"
    hdr[2].text = "Count"
    hdr[3].text = "Status"
    
    # Data rows
    data = [
        ("ULS (total)", "Integrated care units", "39", "Integrated"),
        ("├─ Pre-existing", "Created 1999–2012", "8", "Long-term"),
        ("└─ New (2024)", "Created January 2024", "31", "Newly integrated"),
        ("IPO", "Oncology centers", "3", "Excluded"),
    ]
    
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            row[j].text = cell_text
    
    doc.add_paragraph()
    doc.add_paragraph(
        "The three Portuguese Oncology Institutes (IPO Lisboa, IPO Porto, IPO Coimbra) were excluded from mandatory ULS integration due to their specialized tertiary care mission."
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # RESULTS
    # =========================================================================
    
    doc.add_heading("6. Results", level=1)
    
    doc.add_heading("6.1 Main DiD Estimates", level=2)
    
    # Table: Results
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Estimate"
    hdr[2].text = "P-value"
    
    data = [
        ("Baseline Difference (β₁)", "−0.185", "0.006***"),
        ("DiD Effect (β₃)", "+0.062", "0.433"),
        ("95% CI for β₃", "[−0.093, +0.217]", "—"),
        ("Statistical Power", "25%", "—"),
    ]
    
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            row[j].text = cell_text
    
    doc.add_paragraph()
    
    key_finding = doc.add_paragraph()
    key_finding.add_run("Key Findings:").bold = True
    
    doc.add_paragraph(
        "1. Baseline Difference (β₁ = −0.185, p=0.006): New ULS had 18.5pp lower baseline distress than pre-existing ULS.",
        style='List Number'
    )
    
    doc.add_paragraph(
        "2. DiD Effect (β₃ = +0.062, p=0.433): Not statistically significant. The null finding is uninformative due to severe power limitations.",
        style='List Number'
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # CONCLUSION
    # =========================================================================
    
    doc.add_heading("8. Conclusion", level=1)
    
    doc.add_paragraph(
        "This study provides the first methodological framework and preliminary evidence on Portugal's 2024 universal healthcare integration reform. The main DiD estimate (+6.2pp, 95% CI: −9.3 to +21.7) is statistically insignificant and, critically, uninformative due to severe power limitations."
    )
    
    doc.add_paragraph(
        "The paper's primary contribution is not causal identification—which remains impossible with current data—but establishing evaluation infrastructure and documenting an important selection pattern in early ULS adoption. Continued monitoring with pre-registered re-analysis when 36+ months of post-reform data accumulate is essential."
    )
    
    # =========================================================================
    # SAVE
    # =========================================================================
    
    doc.save("CFE_Empirical_Work_V6_HealthEconomics.docx")
    print("DOCX saved: CFE_Empirical_Work_V6_HealthEconomics.docx")

if __name__ == "__main__":
    create_v6_docx()
